from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\a\Desktop\Work\College\internetAndWeb")
OUTPUT = ROOT / "InternetAndWeb.docx"

QUESTIONS = [
    ("Q01_HTML_Basic_Tags", "Write an HTML code by using basic HTML tags."),
    ("Q02_HTML_Insert_Image", "Write an HTML code to insert image in a webpage."),
    ("Q03_HTML_Hyperlink", "Write an HTML code to implement hyperlink in webpage."),
    ("Q04_HTML_BCA_Timetable", "Write an HTML code to create a timetable of BCA fourth semester."),
    ("Q05_HTML_Lists", "Write an HTML code to demonstrate different types of list."),
    ("Q06_HTML_Employee_Registration_Form", "Write an HTML code to create Employee Registration Form."),
    ("Q07_HTML_CSS_Styles", "Design the webpage by applying inline, internal and external style sheets."),
    ("Q08_HTML_Registration_Form", "Create a Registration Form with the given fields."),
    ("Q09_HTML_JS_Form_Validation", "Write JavaScript to validate the registration form fields."),
    ("Q10_HTML_CSS_JS_Calculator", "Design a simple Calculator by using HTML, CSS and JavaScript."),
    ("Q11_XML_Book_Information", "Write an XML file to display book information."),
    ("Q12_JS_Hello_World", "Write a JavaScript code to print Hello World."),
    ("Q13_JS_Add_Two_Numbers", "Write a JavaScript code to add two numbers entered by user."),
    ("Q14_JS_Even_Odd", "Write a JavaScript code to check if a number is even or odd."),
    ("Q15_JS_Factorial", "Write a JavaScript code to find factorial of a number."),
    ("Q16_JS_Rectangle_Area_Perimeter", "Write a JavaScript code to calculate area and perimeter of a rectangle."),
    ("Q17_JS_Swap_Variables", "Write a JavaScript code to swap two variables."),
    ("Q18_JS_Sum_Natural_Numbers", "Write a JavaScript code to find the sum of natural numbers."),
    ("Q19_JS_Sort_Array_Function", "Write a code to define a user defined function for sorting the values in an array."),
    ("Q20_JS_Fibonacci_Function", "Write a JavaScript code to print the Fibonacci sequence by using functions."),
    ("Q21_JS_Objects_Demo", "Write a code to demonstrate objects in JavaScript."),
    ("Q22_JS_Event_Handling", "Write a code to demonstrate event handling in JavaScript."),
]


def set_cell_border(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_page_layout(document):
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def add_code_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)


def add_index_page(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Internet and Web Technology")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Index")
    run.bold = True
    run.font.size = Pt(12)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.rows[0].cells
    header_cells[0].width = Inches(0.7)
    header_cells[1].width = Inches(6.5)
    header_cells[0].text = "Q.No."
    header_cells[1].text = "Question"

    for cell in header_cells:
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for index, (_, question) in enumerate(QUESTIONS, start=1):
        row = table.add_row().cells
        row[0].width = Inches(0.7)
        row[1].width = Inches(6.5)
        row[0].text = str(index)
        row[1].text = question
        for cell in row:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)


def add_question_page(document, number, folder_name, question):
    document.add_page_break()

    heading = document.add_paragraph()
    run = heading.add_run(f"Question {number}")
    run.bold = True
    run.font.size = Pt(14)

    prompt = document.add_paragraph()
    run = prompt.add_run(question)
    run.font.size = Pt(11)

    folder = ROOT / folder_name
    files = sorted([path for path in folder.iterdir() if path.is_file()])

    for file_path in files:
        label = document.add_paragraph()
        run = label.add_run(f"File: {file_path.name}")
        run.bold = True
        run.font.size = Pt(11)

        content = file_path.read_text(encoding="utf-8")
        for line in content.splitlines():
            add_code_paragraph(document, line)

        if not content.splitlines():
            add_code_paragraph(document, "")


def main():
    document = Document()
    set_page_layout(document)
    add_index_page(document)

    for number, (folder_name, question) in enumerate(QUESTIONS, start=1):
        add_question_page(document, number, folder_name, question)

    document.save(str(OUTPUT))


if __name__ == "__main__":
    main()
